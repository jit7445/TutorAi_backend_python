import os
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from docx import Document

def generate_notes(notes_data: dict, output_dir: str) -> dict:
    os.makedirs(output_dir, exist_ok=True)
    topic = notes_data.get("topic", "Topic")
    notes = notes_data.get("notes", [])
    
    pdf_path = os.path.join(output_dir, "notes.pdf")
    doc_path = os.path.join(output_dir, "notes.docx")
    
    # --- PDF using reportlab (Improved) ---
    doc_pdf = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []
    
    # Add Title
    story.append(Paragraph(f"Topic: {topic}", styles['Title']))
    story.append(Spacer(1, 12))
    
    # Add Notes
    for note in notes:
        story.append(Paragraph(f"• {note}", styles['Normal']))
        story.append(Spacer(1, 6))
        
    doc_pdf.build(story)
    
    # --- DOCX using python-docx ---
    doc_docx = Document()
    doc_docx.add_heading(f"Topic: {topic}", 0)
    
    for note in notes:
        doc_docx.add_paragraph(note, style='List Bullet')
        
    doc_docx.save(doc_path)
    
    return {
        "pdf_path": pdf_path,
        "doc_path": doc_path
    }
